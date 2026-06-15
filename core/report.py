"""
SAKSI — Penyusun Laporan Word (.docx)
=====================================
Membentuk "Laporan Monitoring Harian Transaksi KUPVA BB" mengikuti template
KPwDN Bank Indonesia, lengkap dengan narasi + data + seluruh grafik:

Kop surat: letterhead BI + judul + Periode; baris meta (tanggal cek + cakupan KUPVA).
Seksi 1 : Objek Monitoring & Absensi (jumlah KUPVA, telah/belum upload, donat
          absensi, tabel penyampaian, tabel valuta dilaporkan + deskriptif).
Seksi 2 : Analisis Jumlah Transaksi (growth dtd + grafik volume inline).
Seksi 3 : Analisis Kurs (rasio vs BI + grafik tren kurs & rasio inline).
Seksi 4 : Supervisory Action.
Grafik disisipkan langsung di bawah seksinya agar alur baca mengalir.
Seluruhnya sadar-periode (granularitas pada Konteks).
"""
from __future__ import annotations

from collections import Counter
from io import BytesIO

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import saksi_engine as E

_BIRU_TUA = "002855"
_BIRU_BAR = "1F4E79"
_EMAS = "C8A951"
_BIRU_TEKS = RGBColor(0x1F, 0x4E, 0x79)
_PUTIH = RGBColor(0xFF, 0xFF, 0xFF)


# ----------------------------------------------------------------------------
# Helper docx
# ----------------------------------------------------------------------------
def _shade(element, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def _no_borders(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    table._tbl.tblPr.append(borders)


def _set_cell_bg(cell, fill: str) -> None:
    _shade(cell._tc.get_or_add_tcPr(), fill)


def _banner(doc, judul: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_borders(t)
    cell = t.rows[0].cells[0]
    _set_cell_bg(cell, _BIRU_TUA)
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(judul)
    r1.bold = True
    r1.font.size = Pt(13)
    r1.font.color.rgb = _PUTIH


def _cell_run(cell, text, *, align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, size=10,
              color=_PUTIH, para=None):
    p = para if para is not None else cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def _letterhead(doc, provinsi: str, periode_txt: str, judul: str) -> None:
    """Kop surat gaya Bank Indonesia: logo + BANK INDONESIA (atas), lalu baris
    KPwDN (kiri) · judul laporan (tengah) · Periode (kanan)."""
    t = doc.add_table(rows=2, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_borders(t)

    # Baris 1 — logo + nama bank (satu sel digabung, latar biru tua)
    top = t.rows[0].cells[0].merge(t.rows[0].cells[2])
    _set_cell_bg(top, _BIRU_TUA)
    _cell_run(top, "Ⓑ  BANK INDONESIA", bold=True, size=17)
    _cell_run(top, "BANK SENTRAL REPUBLIK INDONESIA", bold=True, size=8,
              para=top.add_paragraph())

    # Baris 2 — KPwDN | judul | periode (latar biru bar)
    left, mid, right = t.rows[1].cells
    for c in (left, mid, right):
        _set_cell_bg(c, _BIRU_BAR)
    left.width = Inches(2.0)
    mid.width = Inches(3.3)
    right.width = Inches(1.8)
    _cell_run(left, f"KPwDN Provinsi {provinsi}", align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, size=9)
    _cell_run(mid, judul, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11)
    _cell_run(right, f"Periode · {periode_txt}", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, size=9)


_ABU = RGBColor(0x80, 0x80, 0x80)


def _field(paragraph, instr, size=8):
    run = paragraph.add_run()
    run.font.size = Pt(size)
    run.font.color.rgb = _ABU
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(it); run._r.append(e)


def _running(doc, provinsi):
    """Header & footer berulang tiap halaman (judul · KPwDN · Dokumen Internal · Halaman X/Y)."""
    sec = doc.sections[0]

    def _gabu(p, text):
        r = p.add_run(text); r.font.size = Pt(8); r.font.color.rgb = _ABU; return r

    hp = sec.header.paragraphs[0]
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(7.1), WD_TAB_ALIGNMENT.RIGHT)
    _gabu(hp, "Laporan Monitoring Harian Transaksi KUPVA BB")
    _gabu(hp, "\t")
    _gabu(hp, f"KPwDN Provinsi {provinsi}")

    fp = sec.footer.paragraphs[0]
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(7.1), WD_TAB_ALIGNMENT.RIGHT)
    _gabu(fp, "Dokumen Internal - Bank Indonesia")
    _gabu(fp, "\t")
    _gabu(fp, "Halaman ")
    _field(fp, "PAGE")
    _gabu(fp, " dari ")
    _field(fp, "NUMPAGES")


def _center(doc, text, *, bold=False, italic=False, size=10, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color


def _section_bar(doc, text: str) -> None:
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=1)
    _no_borders(t)
    cell = t.rows[0].cells[0]
    _set_cell_bg(cell, _BIRU_BAR)
    r = cell.paragraphs[0].add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = _PUTIH


def _bullet(doc, *segments) -> None:
    p = doc.add_paragraph(style="List Bullet")
    for seg in segments:
        teks, bold = (seg, False) if isinstance(seg, str) else (seg[0], seg[1])
        r = p.add_run(teks)
        r.bold = bold
        r.font.size = Pt(10.5)


def _note(doc, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = _BIRU_TEKS


def _label_grafik(doc, text: str) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = _BIRU_TEKS


def _add_pic(doc, buf: BytesIO, width: float = 6.4) -> None:
    doc.add_picture(buf, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _tabel(doc, headers, rows) -> None:
    """Tabel berisi header berarsir biru + baris data, bergaris (Table Grid)."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        _set_cell_bg(cell, _BIRU_BAR)
        r = cell.paragraphs[0].add_run(str(h))
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = _PUTIH
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            r = cells[j].paragraphs[0].add_run("" if v is None else str(v))
            r.font.size = Pt(9.5)


def _para(doc, *segments, indent: float = 0.25) -> None:
    """Paragraf indent (bukan bullet) dengan run campuran bold/normal."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    for seg in segments:
        teks, bold = (seg, False) if isinstance(seg, str) else (seg[0], seg[1])
        r = p.add_run(teks)
        r.bold = bold
        r.font.size = Pt(10.5)


# ---- format angka gaya laporan asesmen ----
def _pct(x) -> str:
    """Persen rasio (selalu positif), mis. '100.7%'."""
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return "-"
    return f"{x * 100:.1f}%"


def _perub(g) -> str:
    """Arah perubahan tanpa tanda minus: 'naik 16,0%' / 'turun 84,8%' / 'tetap'."""
    if g is None or (isinstance(g, float) and np.isnan(g)):
        return "tidak terdefinisi"
    a = f"{abs(g) * 100:.1f}".replace(".", ",")
    if g > 0:
        return f"naik {a}%"
    if g < 0:
        return f"turun {a}%"
    return "tetap (0,0%)"


def _rp_kata(x) -> str:
    """Rupiah kata penuh, mis. 'Rp12,3 miliar', 'Rp450,0 juta'."""
    x = float(x) if x is not None and x == x else 0.0
    a = abs(x)
    if a >= 1e12:
        return f"Rp{x / 1e12:.1f} triliun".replace(".", ",")
    if a >= 1e9:
        return f"Rp{x / 1e9:.1f} miliar".replace(".", ",")
    if a >= 1e6:
        return f"Rp{x / 1e6:.1f} juta".replace(".", ",")
    return ("Rp" + f"{x:,.0f}").replace(",", ".")


def _stat_kurs(rasio, ambang) -> str:
    """Status kurs 3-tingkat untuk laporan: Normal (≤100%) / Perhatian (>100%) /
    Waspada (≥ ambang, default 105%)."""
    if rasio is None or (isinstance(rasio, float) and np.isnan(rasio)) or rasio == 0:
        return "Tanpa data"
    if rasio >= ambang:
        return "Waspada"
    if rasio > 1.0:
        return "Perhatian"
    return "Normal"


def _stat_vol(g, ambang) -> str:
    """Status volume BERARAH (tidak dinormalkan): Waspada bila NAIK ≥ ambang;
    Perhatian bila TURUN ≥ ambang; selain itu Normal."""
    if g is None or (isinstance(g, float) and np.isnan(g)):
        return "Tanpa data"
    if g >= ambang:
        return "Waspada"
    if g <= -ambang:
        return "Perhatian"
    return "Normal"


# ----------------------------------------------------------------------------
# Grafik (matplotlib → PNG)
# ----------------------------------------------------------------------------
def _png(fig) -> BytesIO:
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=130, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _box(doc, title, *body):
    """Kotak sorot (latar biru muda) berisi judul + paragraf isi."""
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    _set_cell_bg(cell, "EAF1F8")
    rt = cell.paragraphs[0].add_run(title)
    rt.bold = True
    rt.font.size = Pt(10)
    rt.font.color.rgb = _BIRU_TEKS
    for seg in body:
        p = cell.add_paragraph()
        r = p.add_run(seg)
        r.font.size = Pt(10)


def _num(doc, *segments):
    p = doc.add_paragraph(style="List Number")
    for seg in segments:
        teks, bold = (seg, False) if isinstance(seg, str) else (seg[0], seg[1])
        r = p.add_run(teks)
        r.bold = bold
        r.font.size = Pt(10.5)


def _tabel_judul(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = _BIRU_TEKS


_BULAN_PJG = ["Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli",
              "Agustus", "September", "Oktober", "November", "Desember"]


def _tgl_panjang(t) -> str:
    t = pd.Timestamp(t)
    return f"{t.day} {_BULAN_PJG[t.month - 1]} {t.year}"


# ----------------------------------------------------------------------------
# Grafik 3 hari (H-2 s.d. H) — matplotlib, agregat seluruh KUPVA
# ----------------------------------------------------------------------------
def _png(fig) -> BytesIO:
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=130, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


_WK = {"Beli": "#1D9E75", "Tengah": "#185FA5", "Jual": "#E24B4A"}


def _chart_absensi(n_lapor: int, n_belum: int) -> BytesIO:
    fig, ax = plt.subplots(figsize=(3.8, 3.2))
    seg = [(lbl, v, c) for lbl, v, c in
           (("Telah lapor", n_lapor, "#1D9E75"), ("Belum lapor", n_belum, "#E24B4A")) if v > 0]
    if seg:
        ax.pie([s[1] for s in seg], labels=[f"{s[0]}\n({s[1]})" for s in seg],
               colors=[s[2] for s in seg], autopct=lambda p: f"{p:.0f}%",
               startangle=90, textprops=dict(fontsize=8),
               wedgeprops=dict(width=0.42, edgecolor="white"))
    ax.set_title("Absensi penyampaian laporan", fontsize=10)
    ax.axis("equal")
    return _png(fig)


def _chart_kurs3(data, val, pts, hari) -> BytesIO:
    cb = data["combine"]
    xs = [E.fmt_tgl(t) for t in hari]
    rows = {k: [] for k in ("bb", "bt", "bj", "kb", "kt", "kj")}
    for t in hari:
        bi = E.kurs_bi_komponen(data, val, t)
        pt = E.kurs_rata2(cb, t, val, pts, gran="Harian", acuan=bi["tengah"])
        rows["bb"].append(bi["beli"]); rows["bt"].append(bi["tengah"]); rows["bj"].append(bi["jual"])
        rows["kb"].append(pt["beli"]); rows["kt"].append(pt["tengah"]); rows["kj"].append(pt["jual"])
    fig, ax = plt.subplots(figsize=(6.7, 3.0))
    for name, bi_s, k_s in (("Beli", rows["bb"], rows["kb"]),
                            ("Tengah", rows["bt"], rows["kt"]),
                            ("Jual", rows["bj"], rows["kj"])):
        c = _WK[name]
        ax.plot(xs, bi_s, "--s", color=c, ms=4, lw=1.4, label=f"BI {name}")
        ax.plot(xs, k_s, "-o", color=c, ms=5, lw=1.9, label=f"KUPVA {name}")
    ax.set_title(f"Tren kurs {val}: KUPVA BB vs Bank Indonesia (H-2 s.d. H)", fontsize=9.5)
    ax.legend(fontsize=6.5, ncol=3, loc="upper center", bbox_to_anchor=(0.5, -0.16))
    ax.grid(alpha=0.25)
    ax.tick_params(labelsize=8)
    return _png(fig)


def _chart_rasio3(data, val, pts, hari, ambang) -> BytesIO:
    from matplotlib.ticker import FuncFormatter
    cb = data["combine"]
    xs = [E.fmt_tgl(t) for t in hari]
    rb, rt, rj = [], [], []
    for t in hari:
        bi = E.kurs_bi_komponen(data, val, t)
        pt = E.kurs_rata2(cb, t, val, pts, gran="Harian", acuan=bi["tengah"])
        rb.append(E.hitung_rasio(pt["beli"], bi["beli"]))
        rt.append(E.hitung_rasio(pt["tengah"], bi["tengah"]))
        rj.append(E.hitung_rasio(pt["jual"], bi["jual"]))
    fig, ax = plt.subplots(figsize=(6.7, 3.0))
    ax.plot(xs, rb, "-o", color=_WK["Beli"], lw=1.9, label="Rasio Beli")
    ax.plot(xs, rt, "-o", color=_WK["Tengah"], lw=1.9, label="Rasio Tengah")
    ax.plot(xs, rj, "-o", color=_WK["Jual"], lw=1.9, label="Rasio Jual")
    ax.axhline(ambang, color="#E24B4A", ls="--", lw=1, label=f"Waspada {ambang:.0%}")
    ax.axhline(1.0, color="#BA7517", ls=":", lw=1, label="100%")
    ax.yaxis.set_major_formatter(FuncFormatter(lambda y, _: f"{y * 100:.1f}%"))
    ax.set_title(f"Tren rasio kurs {val} terhadap Bank Indonesia (H-2 s.d. H)", fontsize=9.5)
    ax.legend(fontsize=6.5, ncol=3, loc="upper center", bbox_to_anchor=(0.5, -0.16))
    ax.grid(alpha=0.25)
    ax.tick_params(labelsize=8)
    return _png(fig)


def _chart_volume3(data, val, pts, hari) -> BytesIO:
    cb = data["combine"]
    xs = [E.fmt_tgl(t) for t in hari]
    J, B = [], []
    for t in hari:
        j, b = E.volume_jual_beli_lapor(cb, t, [val], pts, gran="Harian")
        J.append(j); B.append(b)
    x = np.arange(len(xs))
    fig, ax = plt.subplots(figsize=(6.7, 3.0))
    ax.bar(x - 0.2, J, 0.4, color="#185FA5", label="Volume Jual")
    ax.bar(x + 0.2, B, 0.4, color="#1D9E75", label="Volume Beli")
    ax.set_xticks(x)
    ax.set_xticklabels(xs, fontsize=8)
    ax.set_title(f"Tren volume transaksi {val} (Jual & Beli, Rp) — H-2 s.d. H", fontsize=9.5)
    ax.legend(fontsize=7)
    ax.grid(alpha=0.25, axis="y")
    ax.tick_params(axis="y", labelsize=8)
    return _png(fig)


def _chart_growth3(data, val, pts, hari_all, hari, ambang) -> BytesIO:
    from matplotlib.ticker import FuncFormatter
    cb = data["combine"]
    xs = [E.fmt_tgl(t) for t in hari]
    GJ, GB = [], []
    for t in hari:
        it = hari_all.index(t)
        prev = hari_all[it - 1] if it > 0 else None
        jh, bh = E.volume_jual_beli_lapor(cb, t, [val], pts, gran="Harian")
        if prev is not None:
            jp, bp = E.volume_jual_beli_lapor(cb, prev, [val], pts, gran="Harian")
        else:
            jp = bp = np.nan
        GJ.append(E.growth(jh, jp)); GB.append(E.growth(bh, bp))
    fig, ax = plt.subplots(figsize=(6.7, 3.0))
    ax.plot(xs, GJ, "-o", color="#185FA5", lw=1.9, label="Growth Jual")
    ax.plot(xs, GB, "-o", color="#1D9E75", lw=1.9, label="Growth Beli")
    ax.axhline(ambang, color="#E24B4A", ls="--", lw=1)
    ax.axhline(-ambang, color="#E24B4A", ls="--", lw=1)
    ax.axhline(0, color="#888780", ls=":", lw=1)
    ax.yaxis.set_major_formatter(FuncFormatter(lambda y, _: f"{y * 100:.0f}%"))
    ax.set_title(f"Tren pertumbuhan dtd {val} (Jual & Beli) — H-2 s.d. H", fontsize=9.5)
    ax.legend(fontsize=7)
    ax.grid(alpha=0.25)
    ax.tick_params(labelsize=8)
    return _png(fig)


# ----------------------------------------------------------------------------
# Penyusun utama
# ----------------------------------------------------------------------------
def build_report(ctx, provinsi: str = "DKI Jakarta", kota: str = "",
                 pengesahan=None) -> bytes:
    """Bangun laporan .docx (format 4 halaman) mengikuti template SAKSI.

    `pengesahan` = daftar 4 tuple (header, peran, nama, jabatan) untuk Lembar
    Pengesahan. Bila None, dipakai default (header: Dipersiapkan/Diperiksa/
    Didukung/Disetujui oleh) dengan nama kosong."""
    g = ctx.granularitas
    data, cb = ctx.data, ctx.data["combine"]
    pts_all = E.daftar_pt(cb)
    total = len(pts_all)
    tgl_h, tgl_p = ctx.tgl_h, ctx.tgl_p
    ar, av = ctx.ambang_rasio, ctx.ambang_dtd
    kota = kota or provinsi
    hari_all = E.daftar_tanggal(cb)
    hari = [t for t in hari_all if pd.Timestamp(t) <= pd.Timestamp(tgl_h)][-3:]

    # ---- Absensi (§3) ----
    absn = E.tabel_absensi(data, tgl_h, pts_all, gran=g)
    n_lapor = int((absn["Status"] == "Lengkap").sum())
    belum = absn[absn["Status"] == "Belum lapor"]["KUPVA BB"].tolist()
    n_belum = len(belum)
    ket = (n_lapor / total * 100) if total else 0.0

    # ---- Ringkasan transaksi ----
    sub_h = E.filter_cb(cb, tgl=tgl_h, gran=g)
    v_jual = float(sub_h[E.C_JUAL_RP].sum())
    v_beli = float(sub_h[E.C_BELI_RP].sum())

    # ---- Matriks terintegrasi ----
    mtx = E.ringkasan_kupva(data, tgl_h, tgl_p, ar, av, gran=g)
    n_wk = int((mtx["Status Kurs"] == "Waspada").sum())
    n_wv = int((mtx["Status Volume"] == "Waspada").sum())
    daftar_wk = mtx[mtx["Status Kurs"] == "Waspada"]["KUPVA BB"].tolist()
    daftar_wv = mtx[mtx["Status Volume"] == "Waspada"]["KUPVA BB"].tolist()

    # ---- Valuta dilaporkan ----
    vals_h = E.valuta_pt_pada(data, pts_all, tgl_h, gran=g)
    val = ctx.valuta_fokus if ctx.valuta_fokus in vals_h else (vals_h[0] if vals_h else ctx.valuta_fokus)
    cnt = Counter()
    for _, r in absn[absn["Status"] == "Lengkap"].iterrows():
        cnt.update(E.valuta_pt_pada(data, r["ID"], tgl_h, gran=g))
    top = cnt.most_common()
    sebaran = "; ".join(f"{v} oleh {n} KUPVA BB" for v, n in top) if top else "-"

    # ---- Kurs ilustrasi (agregat) ----
    bi_h = E.kurs_bi_komponen(data, val, tgl_h)
    pt_h = E.kurs_rata2(cb, tgl_h, val, pts_all, gran=g, acuan=bi_h["tengah"])
    r_beli = E.hitung_rasio(pt_h["beli"], bi_h["beli"])
    r_jual = E.hitung_rasio(pt_h["jual"], bi_h["jual"])
    r_teng = E.hitung_rasio(pt_h["tengah"], bi_h["tengah"])

    # ---- Volume agregat ----
    jp = E.filter_cb(cb, tgl=tgl_p, gran=g)[E.C_JUAL_RP].sum() if tgl_p is not None else np.nan
    bp = E.filter_cb(cb, tgl=tgl_p, gran=g)[E.C_BELI_RP].sum() if tgl_p is not None else np.nan
    g_jual = E.growth(v_jual, jp)
    g_beli = E.growth(v_beli, bp)

    # ---- Rincian KUPVA perubahan signifikan ----
    rinci = []
    for _, r in mtx[mtx["Status Volume"] == "Waspada"].iterrows():
        pid = r["ID"]
        jh, bh = E.volume_jual_beli_lapor(cb, tgl_h, None, [pid], gran=g)
        j2, b2 = (E.volume_jual_beli_lapor(cb, tgl_p, None, [pid], gran=g)
                  if tgl_p is not None else (np.nan, np.nan))
        rinci.append(f"{r['KUPVA BB']} (vol jual {_perub(E.growth(jh, j2))} setara {_rp_kata(jh)}; "
                     f"vol beli {_perub(E.growth(bh, b2))} setara {_rp_kata(bh)})")
    rincian_txt = "; ".join(rinci)

    pbd = E.fmt_tgl(tgl_p) if tgl_p is not None else "-"

    # ---- dokumen ----
    doc = Document()
    for sct in doc.sections:
        sct.top_margin = Inches(0.5)
        sct.bottom_margin = Inches(0.5)
        sct.left_margin = Inches(0.7)
        sct.right_margin = Inches(0.7)
    _running(doc, provinsi)

    _letterhead(doc, provinsi, pd.Timestamp(tgl_h).strftime("%d/%m/%Y"),
                f"Laporan Monitoring {g} Transaksi KUPVA BB")
    _center(doc, f"Tanggal cek (H): {_tgl_panjang(tgl_h)}   ·   Pembanding (H-1): {pbd}   ·   "
                 f"Objek monitoring: {total} KUPVA BB", size=9)

    # ====================== BAGIAN 1 — ABSENSI ======================
    _section_bar(doc, "BAGIAN 1 · OBJEK MONITORING DAN ABSENSI LAPORAN")
    _note(doc, "Penyampaian, ketepatan waktu (H+1 pukul 12.00), dan kelengkapan laporan - modul Absensi (§3).")
    _para(doc, "Aspek absensi merupakan fondasi dari keseluruhan asesmen harian, karena keandalan "
               "analisis kurs (Bagian 2) maupun jumlah transaksi (Bagian 3) sepenuhnya bergantung pada "
               "kelengkapan dan ketepatan waktu laporan yang disampaikan KUPVA BB. Pada tanggal cek "
               f"{_tgl_panjang(tgl_h)}, objek monitoring mencakup ", (f"{total} KUPVA BB", True),
               ". Objek tersebut dipilih karena secara kumulatif merepresentasikan mayoritas - minimum "
               f"50% - nilai transaksi jual dan beli di wilayah kerja KPwDN Provinsi {provinsi}, sehingga "
               "gambaran yang dihasilkan telah cukup mewakili dinamika transaksi penukaran valuta asing "
               "di wilayah dimaksud.", indent=0)
    _para(doc, f"Dari {total} objek monitoring tersebut, sebanyak ",
          (f"{n_lapor} KUPVA BB telah menyampaikan", True), " (mengunggah) laporan transaksi harian, "
          "sedangkan ", (f"{n_belum} KUPVA BB belum menyampaikan", True),
          f", sehingga ketepatan penyampaian tercatat ", (f"{ket:.1f}%", True),
          ". Penilaian mencakup dua dimensi yang berbeda namun saling melengkapi, yaitu kelengkapan - "
          "dinilai dari ketersediaan catatan transaksi pada hari pelaporan - dan ketepatan waktu "
          "terhadap batas H+1 maksimal pukul 12.00 waktu setempat. Mengingat data sumber belum memuat "
          "cap waktu (timestamp) penyampaian, penilaian ketepatan waktu dilakukan secara manual oleh "
          "KPwDN. Tingkat ketepatan penyampaian merupakan indikator awal kepatuhan pelaporan: ketepatan "
          "yang menurun maupun KUPVA BB yang berulang kali tidak menyampaikan laporan perlu dicermati "
          "sebagai potensi titik buta (blind spot) pengawasan sekaligus indikasi persoalan kepatuhan "
          "yang lebih luas.", indent=0)
    _add_pic(doc, _chart_absensi(n_lapor, n_belum), width=3.4)

    _tabel_judul(doc, "Tabel 1.1 - Status penyampaian laporan per KUPVA BB")
    rows11 = [[r["KUPVA BB"], "Sudah upload" if r["Status"] == "Lengkap" else "Belum upload",
               len(E.valuta_pt_pada(data, r["ID"], tgl_h, gran=g)), E.rupiah(r["Volume H (Rp)"])]
              for _, r in absn.iterrows()]
    _tabel(doc, ["KUPVA BB", "Penyampaian", "Jml valuta", "Volume (Rp)"], rows11)

    _box(doc, "KUPVA BB yang perlu ditindaklanjuti",
         f"Belum menyampaikan laporan pada tanggal cek: {', '.join(belum) if belum else 'tidak ada'}.",
         "Tindak lanjut atas kelengkapan dan ketepatan waktu penyampaian diuraikan pada Bagian 4 "
         "(Supervisory Action).")

    _tabel_judul(doc, "Tabel 1.2 - Jenis valuta yang dilaporkan per KUPVA BB")
    rows12 = []
    for _, r in absn[absn["Status"] == "Lengkap"].iterrows():
        vs = E.valuta_pt_pada(data, r["ID"], tgl_h, gran=g)
        rows12.append([r["KUPVA BB"], ", ".join(vs) if vs else "-", len(vs)])
    _tabel(doc, ["KUPVA BB", "Valuta dilaporkan", "Jml valuta"], rows12)
    if top:
        _para(doc, f"Pada aspek kelengkapan jenis valuta, terdapat ", (f"{len(cnt)} jenis valuta", True),
              f" yang dilaporkan oleh {n_lapor} KUPVA BB pada tanggal cek. Valuta yang paling banyak "
              f"dilaporkan adalah ", (f"{top[0][0]} (oleh {top[0][1]} dari {n_lapor} KUPVA BB)", True),
              ". Konsentrasi pelaporan pada valuta tertentu menggambarkan profil dan preferensi "
              "transaksi nasabah di wilayah kerja, sekaligus menjadi penanda valuta mana yang paling "
              f"perlu dicermati pada analisis kurs (Bagian 2) dan volume (Bagian 3). Sebaran "
              f"selengkapnya: {sebaran}.", indent=0)
    _note(doc, "Kriteria penilaian: KUPVA BB dinilai telah menyampaikan apabila terdapat catatan "
               "transaksi pada hari pelaporan. Oleh karena data sumber tidak memuat cap waktu "
               "penyampaian, penilaian ketepatan terhadap batas H+1 pukul 12.00 dilakukan secara manual "
               "oleh KPwDN.")

    # ====================== BAGIAN 2 — KURS ======================
    _section_bar(doc, f"BAGIAN 2 · ANALISIS MONITORING TRANSAKSI {g.upper()} — KURS")
    _note(doc, "Kewajaran kurs KUPVA BB terhadap kurs acuan Bank Indonesia - modul Kurs (§1).")
    _para(doc, "Analisis kurs menilai kewajaran kurs yang ditetapkan KUPVA BB dengan membandingkannya "
               "terhadap kurs acuan Bank Indonesia untuk masing-masing jenis valuta, melalui rasio "
               "kurs = kurs KUPVA ÷ kurs Bank Indonesia. Pemantauan ini penting untuk menjaga "
               "perlindungan konsumen dan integritas pasar, sebab deviasi yang terlalu lebar terhadap "
               "acuan dapat mengindikasikan margin yang berlebihan atau praktik penetapan kurs yang "
               "tidak wajar. Rasio dikategorikan ke dalam tiga tingkat, yaitu Normal apabila tidak "
               "melebihi 100%, Perhatian apabila berada di atas 100%, dan Waspada apabila mencapai "
               f"{ar:.0%} atau lebih terhadap kurs acuan Bank Indonesia. KPwDN menghitung rata-rata "
               "rasio atas seluruh objek monitoring untuk tiap jenis valuta; rasio di atas 100% menjadi "
               f"perhatian pengawasan, dengan penekanan khusus pada rasio yang menyentuh ambang Waspada "
               f"(≥ {ar:.0%}) karena menandakan deviasi yang signifikan.", indent=0)
    _para(doc, "Rasio dirinci ke dalam tiga komponen - kurs beli, kurs jual, dan kurs tengah - yang "
               "masing-masing memberikan informasi berbeda. Status akhir kurs per valuta ditetapkan "
               "dari kategori terburuk di antara ketiga komponen tersebut, sehingga satu komponen yang "
               "menyentuh ambang Waspada sudah cukup menjadikan valuta dimaksud berstatus Waspada. "
               f"Sebagai ilustrasi untuk valuta {val}, ",
          (f"rasio kurs beli tercatat {_pct(r_beli)} ({_stat_kurs(r_beli, ar)}), kurs jual {_pct(r_jual)} "
           f"({_stat_kurs(r_jual, ar)}), dan kurs tengah {_pct(r_teng)} ({_stat_kurs(r_teng, ar)})", True),
          ". Selain besaran pada tanggal cek, arah dan persistensi pergerakan juga penting untuk "
          "dicermati: deviasi yang muncul sesaat memiliki bobot pengawasan yang berbeda dengan deviasi "
          "yang berulang selama beberapa hari. Tren kurs dan tren rasio terhadap acuan sepanjang H-2 "
          "sampai dengan H pelaporan disajikan pada grafik berikut, sedangkan rincian rasio per jenis "
          "valuta atas agregat seluruh KUPVA BB disajikan pada Tabel 2.1.", indent=0)
    _tabel_judul(doc, f"Grafik 2.1 - Tren kurs {val} (Beli · Tengah · Jual): KUPVA BB vs Bank Indonesia · H-2 s.d. H")
    _add_pic(doc, _chart_kurs3(data, val, pts_all, hari))
    _tabel_judul(doc, f"Grafik 2.2 - Tren rasio kurs {val} (Beli · Tengah · Jual) terhadap Bank Indonesia · H-2 s.d. H")
    _add_pic(doc, _chart_rasio3(data, val, pts_all, hari, ar))

    _tabel_judul(doc, "Tabel 2.1 - Rincian rasio kurs per valuta (agregat seluruh KUPVA BB)")
    tk = E.tabel_kurs_komponen(data, pts_all, tgl_h, vals_h, ar, gran=g)
    rows21 = [[r["Valuta"], _pct(r["Rasio Beli"]), _pct(r["Rasio Jual"]), _pct(r["Rasio Tengah"]),
               r["Status Akhir"]] for _, r in tk.iterrows()]
    _tabel(doc, ["Valuta", "Rasio Beli", "Rasio Jual", "Rasio Tengah", "Status Akhir"], rows21)

    if daftar_wk:
        narasi_wk = (f"Terdapat {n_wk} KUPVA BB dengan rasio kurs berkategori Waspada (≥ {ar:.0%}) pada "
                     f"tanggal cek, yaitu {', '.join(daftar_wk)}.")
    else:
        narasi_wk = (f"Tidak terdapat KUPVA BB dengan rasio kurs berkategori Waspada (≥ {ar:.0%}) pada "
                     "tanggal cek; seluruh objek monitoring tergolong Normal/Perhatian. Rasio di atas "
                     "100% tetap menjadi perhatian pengawasan.")
    _box(doc, "Temuan dan pendalaman — Kurs", narasi_wk,
         "Pendalaman penyebab terhadap KUPVA BB berkategori Waspada dilakukan dengan menginformasikan "
         "nama KUPVA BB, jenis valuta, serta kurs (beli/tengah/jual) yang digunakan, sebagaimana "
         "ditindaklanjuti pada Bagian 4.")

    # ====================== BAGIAN 3 — VOLUME ======================
    _section_bar(doc, f"BAGIAN 3 · ANALISIS MONITORING TRANSAKSI {g.upper()} — JUMLAH TRANSAKSI")
    _note(doc, "Pertumbuhan volume transaksi day-to-day (dtd) - modul Volume (§2).")
    _para(doc, "Analisis jumlah transaksi memantau pergerakan volume transaksi harian guna mendeteksi "
               "anomali yang berpotensi mengindikasikan transaksi tidak wajar. Pengukuran dilakukan "
               "melalui pertumbuhan day-to-day (dtd) = (nilai tanggal cek - nilai pembanding) ÷ nilai "
               f"pembanding, yang membandingkan {E.fmt_tgl(tgl_h)} terhadap tanggal transaksi pembanding "
               f"{pbd}. Suatu transaksi dikategorikan Waspada apabila terjadi perubahan - baik "
               f"peningkatan maupun penurunan - sebesar {av:.0%} atau lebih (dtd) pada sisi jual dan/atau "
               f"beli, dan Normal apabila perubahan di bawah {av:.0%}. Lonjakan naik (≥ {av:.0%}) dapat "
               f"mengindikasikan transaksi yang tidak biasa, sedangkan penurunan tajam (turun ≥ {av:.0%}) "
               "dapat mengindikasikan pergeseran aktivitas ke kanal lain ataupun underreporting; keduanya "
               "sama-sama relevan dari perspektif kewaspadaan dan pencegahan APU-PPT.", indent=0)
    _para(doc, "Secara agregat untuk seluruh valuta, ",
          (f"volume jual {_perub(g_jual)} dtd menjadi setara {_rp_kata(v_jual)} ({_stat_vol(g_jual, av)}) "
           f"dan volume beli {_perub(g_beli)} dtd menjadi setara {_rp_kata(v_beli)} "
           f"({_stat_vol(g_beli, av)})", True),
          ". Selisih ini menggambarkan posisi neto transaksi pada tanggal cek; ketidakseimbangan yang "
          "besar antara sisi jual dan beli dapat menjadi bahan pencermatan lebih lanjut. Status akhir "
          "volume per valuta ditetapkan dari kategori terburuk antar sisi (jual/beli), sehingga anomali "
          "pada satu sisi sudah cukup memunculkan status Waspada. Tren volume transaksi (jual dan beli) "
          "serta pertumbuhan dtd sepanjang H-2 sampai dengan H pelaporan disajikan pada grafik berikut, "
          "sedangkan rincian pertumbuhan per jenis valuta atas agregat seluruh KUPVA BB disajikan pada "
          "Tabel 3.1.", indent=0)
    _tabel_judul(doc, f"Grafik 3.1 - Tren volume transaksi {val} (Jual · Beli) · H-2 s.d. H")
    _add_pic(doc, _chart_volume3(data, val, pts_all, hari))
    _tabel_judul(doc, f"Grafik 3.2 - Tren pertumbuhan dtd {val} (Jual · Beli) · H-2 s.d. H")
    _add_pic(doc, _chart_growth3(data, val, pts_all, hari_all, hari, av))

    _tabel_judul(doc, "Tabel 3.1 - Rincian volume dan pertumbuhan per valuta (agregat seluruh KUPVA BB)")
    tv = E.tabel_volume_komponen(data, pts_all, tgl_h, tgl_p, vals_h, av, gran=g)
    rows31 = [[r["Valuta"], E.rupiah(r["Jual (H)"]), _perub(r["Growth Jual"]),
               E.rupiah(r["Beli (H)"]), _perub(r["Growth Beli"]), r["Status Akhir"]]
              for _, r in tv.iterrows()]
    _tabel(doc, ["Valuta", "Vol Jual (H)", "Growth Jual", "Vol Beli (H)", "Growth Beli", "Status Akhir"], rows31)

    body3 = []
    if daftar_wv:
        body3.append(f"Terdapat {n_wv} KUPVA BB dengan jumlah transaksi berkategori Waspada "
                     f"(perubahan dtd ≥ {av:.0%} pada sisi jual dan/atau beli), yaitu: {', '.join(daftar_wv)}.")
    else:
        body3.append(f"Tidak terdapat KUPVA BB dengan perubahan jumlah transaksi mencapai {av:.0%} atau "
                     "lebih (kategori Waspada) pada tanggal cek.")
    if rincian_txt:
        body3.append(f"Rincian KUPVA BB dengan perubahan volume signifikan: {rincian_txt}.")
    body3.append("Pendalaman penyebab dilakukan dengan menginformasikan nama KUPVA BB, jenis valuta, "
                 "serta kurs (beli/tengah/jual) yang digunakan, sebagaimana ditindaklanjuti pada Bagian 4.")
    _box(doc, "Temuan dan pendalaman — Jumlah transaksi", *body3)

    # ====================== BAGIAN 4 — SUPERVISORY ======================
    _section_bar(doc, "BAGIAN 4 · SUPERVISORY ACTION")
    _note(doc, "Sintesa status akhir per KUPVA BB dan rekomendasi tindakan pengawasan.")
    _para(doc, "Hasil analisis absensi (Bagian 1), kurs (Bagian 2), dan jumlah transaksi (Bagian 3) "
               "disintesiskan menjadi status akhir terintegrasi per KUPVA BB sebagaimana disajikan pada "
               "Tabel 4.1. Status akhir ditetapkan dari gabungan kondisi terburuk antar aspek kurs dan "
               "aspek volume - sehingga status Waspada pada salah satu aspek akan menjadikan status "
               "akhir KUPVA BB tersebut Waspada - sedangkan KUPVA BB yang belum menyampaikan laporan "
               "berstatus “Tanpa data”. Status akhir inilah yang menentukan bentuk dan intensitas "
               "tindakan pengawasan: kategori Normal ditangani melalui pemantauan offsite rutin, "
               "kategori Perhatian melalui pemantauan yang lebih dekat, kategori Waspada melalui "
               "pendalaman dan klarifikasi, sedangkan status Tanpa data ditindaklanjuti melalui "
               "penegakan kewajiban pelaporan.", indent=0)
    _tabel_judul(doc, "Tabel 4.1 - Matriks terintegrasi dan status akhir per KUPVA BB")
    rows41 = [[r["KUPVA BB"], r["Absensi"], r["Status Kurs"], r["Status Volume"], r["Status Akhir"]]
              for _, r in mtx.iterrows()]
    _tabel(doc, ["KUPVA BB", "Absensi", "Status Kurs", "Status Volume", "Status Akhir"], rows41)
    _para(doc, "Mengacu pada hasil di atas, KPwDN menetapkan tindakan pengawasan sebagai berikut.", indent=0)
    _num(doc, ("Pengawasan offsite atas KUPVA BB kategori Normal. ", True),
         "Melaksanakan pemantauan offsite secara berkelanjutan terhadap KUPVA BB berkategori Normal, "
         "dengan intensitas yang ditingkatkan pada periode terjadinya gejolak nilai tukar, guna "
         "memastikan kurs dan volume transaksi tetap berada pada kisaran wajar tanpa memerlukan "
         "tindakan korektif lebih lanjut.")
    _num(doc, ("Tindak lanjut penyampaian laporan bagi KUPVA BB belum/sebagian lapor. ", True),
         f"Menyampaikan tindak lanjut kepada {', '.join(belum) if belum else 'tidak ada'} atas "
         "kelengkapan dan ketepatan waktu penyampaian laporan transaksi harian sesuai batas H+1 pukul "
         "12.00 waktu setempat, serta memantau kepatuhan penyampaian pada periode berikutnya.")
    _num(doc, ("Pendalaman atas KUPVA BB kategori Waspada. ", True),
         "Melakukan pendalaman penyebab terhadap KUPVA BB dengan kurs berkategori Waspada "
         f"({', '.join(daftar_wk) if daftar_wk else 'tidak ada'}) dan/atau jumlah transaksi berkategori "
         f"Waspada ({', '.join(daftar_wv) if daftar_wv else 'tidak ada'}). Pendalaman dilakukan dengan "
         "menginformasikan nama KUPVA BB, jenis valuta, serta kurs (beli/tengah/jual) yang digunakan, "
         "dan apabila diperlukan ditindaklanjuti dengan klarifikasi langsung atau permintaan penjelasan "
         "kepada KUPVA BB terkait.")
    _para(doc, "Seluruh tindakan pengawasan di atas didokumentasikan sebagai dasar pemantauan harian "
               "yang berkesinambungan. Eskalasi tindak lanjut dilakukan apabila kondisi Waspada bersifat "
               "persisten atau berulang pada KUPVA BB yang sama, atau apabila pendalaman mengindikasikan "
               "adanya potensi pelanggaran ketentuan.", indent=0)

    # ====================== LEMBAR PENGESAHAN (4 kolom) ======================
    _section_bar(doc, "LEMBAR PENGESAHAN")
    _para(doc, f"{kota}, {_tgl_panjang(tgl_h)}", indent=0)
    if pengesahan is None:
        pengesahan = [("Dipersiapkan oleh", "Pelaksana", "", "Staf"),
                      ("Diperiksa oleh", "Pengawas Senior", "", "Asisten Direktur"),
                      ("Didukung oleh", "Pengawas Senior", "", "Asisten Direktur"),
                      ("Disetujui oleh", "Pengawas Eksekutif", "", "Deputi Direktur")]
    ts = doc.add_table(rows=2, cols=len(pengesahan))
    ts.style = "Table Grid"
    ts.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (hdr, peran, nama, jab) in enumerate(pengesahan):
        hc = ts.rows[0].cells[j]
        _set_cell_bg(hc, _BIRU_BAR)
        _cell_run(hc, hdr, bold=True, size=10)
        _cell_run(hc, peran, size=8, para=hc.add_paragraph())
        bc = ts.rows[1].cells[j]
        _cell_run(bc, "\n\n\n" + (nama or "(..............................)"),
                  bold=True, size=10, color=RGBColor(0, 0, 0))
        _cell_run(bc, jab, size=9, color=RGBColor(0, 0, 0), para=bc.add_paragraph())

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
